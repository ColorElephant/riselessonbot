# app.py
import os, tempfile, json, time
from flask import Flask, request, jsonify
import requests
from docx import Document
from duckduckgo_search import ddg
from readability import Document as ReadabilityDoc
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader

# Optional OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# CONFIG
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
DEFAULT_TEMPLATE_PATH = os.environ.get("DEFAULT_TEMPLATE_PATH", "/mnt/data/Sample Lesson Plan.docx")
# If you want to change default path, set DEFAULT_TEMPLATE_PATH env var in Render

if not TELEGRAM_TOKEN:
    raise RuntimeError("Set TELEGRAM_TOKEN env var")

BASE_TELEGRAM_URL = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}"
app = Flask(__name__)

# In-memory sessions (ephemeral)
SESS = {}  # chat_id -> state dict

def telegram_api(method, params=None, files=None, json_payload=None):
    url = f"{BASE_TELEGRAM_URL}/{method}"
    if files:
        return requests.post(url, params=params, files=files, timeout=30)
    if json_payload:
        return requests.post(url, json=json_payload, timeout=30)
    return requests.post(url, data=params, timeout=30)

# Helpers
def send_message(chat_id, text, reply_markup=None):
    payload = {"chat_id": chat_id, "text": text}
    if reply_markup:
        payload["reply_markup"] = json.dumps(reply_markup)
    telegram_api("sendMessage", params=payload)

def download_file(file_id, dest_path):
    r = telegram_api("getFile", params={"file_id": file_id})
    r.raise_for_status()
    data = r.json()
    file_path = data["result"]["file_path"]
    file_url = f"https://api.telegram.org/file/bot{TELEGRAM_TOKEN}/{file_path}"
    r2 = requests.get(file_url, timeout=60)
    r2.raise_for_status()
    with open(dest_path, "wb") as f:
        f.write(r2.content)
    return dest_path

def extract_text_from_pdf(path):
    text = []
    reader = PdfReader(path)
    for p in reader.pages:
        t = p.extract_text()
        if t:
            text.append(t)
    return "\n".join(text)

def extract_text_from_url(url, max_chars=20000):
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent":"Mozilla/5.0"})
        r.raise_for_status()
    except Exception as e:
        return ""
    # Use readability to get main content
    doc = ReadabilityDoc(r.text)
    content_html = doc.summary()
    soup = BeautifulSoup(content_html, "html.parser")
    text = soup.get_text(separator="\n")
    return text[:max_chars]

def summarize_text(text, sentences_count=6):
    # fallback: simple substring if text tiny
    if not text:
        return ""
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = TextRankSummarizer()
    summary_sentences = summarizer(parser.document, sentences_count)
    return "\n".join(str(s) for s in summary_sentences)

def extract_objectives_from_text(text, max_points=5):
    # crude heuristic: pick top sentences with 'able to', 'will', 'understand', or use summary of key sentences
    lines = []
    lowered = text.lower()
    candidates = []
    for sent in text.split("."):
        s = sent.strip()
        if not s:
            continue
        sl = s.lower()
        if any(k in sl for k in ("able to", "will", "understand", "learn", "identify", "describe")):
            candidates.append(s.strip())
    if candidates:
        return "\n".join(f"• {c}" for c in candidates[:max_points])
    # fallback: take first N summary sentences
    summ = summarize_text(text, sentences_count=max_points)
    if summ:
        return "\n".join(f"• {s.strip()}" for s in summ.split("\n") if s.strip())
    return "• Objective 1\n• Objective 2"

def generate_activities(text, max_items=4):
    # simple template: comprehension, group activity, experiment/demo, assessment
    return (
        "1. Read the summary and discuss key terms.\n"
        "2. Small-group activity: identify examples from the text.\n"
        "3. Hands-on/demo (if applicable): follow the experiment steps.\n"
        "4. Exit ticket: one short question to assess learning."
    )

def generate_assessment_questions(text, max_q=4):
    # pick top sentences and convert to short questions (heuristic)
    summary = summarize_text(text, sentences_count=4)
    qs = []
    for i, s in enumerate(summary.split("\n")[:max_q]):
        s = s.strip().rstrip(".")
        if len(s.split()) < 3:
            continue
        qs.append(f"Q{i+1}. Explain: {s}?")
    if not qs:
        qs = ["Q1. What is the main idea of the chapter?", "Q2. List two key points."]
    return "\n".join(qs)

def fill_template_and_send(chat_id, title, summary, objectives, activities, assessment, references=""):
    template_path = SESS.get(chat_id, {}).get("template_path") or DEFAULT_TEMPLATE_PATH
    if not os.path.exists(template_path):
        send_message(chat_id, "Template not found on server. Please upload a .docx template or set DEFAULT_TEMPLATE_PATH.")
        return
    # prepare replacements (adjust keys as your template has)
    replacements = {
        "{{ChapterTitle}}": title,
        "{{Summary}}": summary,
        "{{LearningObjectives}}": objectives,
        "{{Activities}}": activities,
        "{{Assessment}}": assessment,
        "{{References}}": references
    }
    out_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    out_path = out_tmp.name
    doc = Document(template_path)
    # simple replace
    for p in doc.paragraphs:
        for k, v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in replacements.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)
    doc.save(out_path)
    # send file
    with open(out_path, "rb") as f:
        files = {"document": (os.path.basename(out_path), f)}
        telegram_api("sendDocument", params={"chat_id": chat_id}, files=files)
    send_message(chat_id, "Lesson plan generated ✅")

# Webhook routes
@app.route("/health", methods=["GET"])
def health():
    return jsonify({"ok": True})

@app.route("/webhook", methods=["POST"])
def webhook():
    update = request.get_json(force=True)
    if not update or "message" not in update:
        return jsonify({"ok": True})
    msg = update["message"]
    chat_id = msg["chat"]["id"]
    SESS.setdefault(chat_id, {})
    # simple command to start
    if "text" in msg and msg["text"].strip() == "/Hi_Rise":
        # keyboard with options
        kb = {"keyboard":[["Upload PDF"],["Scan Text (Image)"],["Paste Text"],["Ask Bot to Find Lesson"]],"one_time_keyboard":True,"resize_keyboard":True}
        send_message(chat_id, "Hi — for which lesson shall we create a lesson plan? Choose how to provide the lesson:", reply_markup=kb)
        SESS[chat_id]["state"] = "idle"
        return jsonify({"ok": True})

    # handle document uploads (PDF or template)
    if "document" in msg:
        doc = msg["document"]
        fname = doc.get("file_name","file")
        fid = doc["file_id"]
        tmpdir = tempfile.mkdtemp()
        local_path = os.path.join(tmpdir, fname)
        try:
            download_file(fid, local_path)
        except Exception as e:
            send_message(chat_id, f"Failed to download file: {e}")
            return jsonify({"ok": True})
        # if .docx treat as custom template
        if fname.lower().endswith(".docx"):
            SESS[chat_id]["template_path"] = local_path
            send_message(chat_id, "Template saved for your session. Now upload or paste the lesson content or use /Hi_Rise again.")
            return jsonify({"ok": True})
        # if pdf treat as chapter pdf
        if fname.lower().endswith(".pdf"):
            try:
                text = extract_text_from_pdf(local_path)
            except Exception as e:
                send_message(chat_id, f"PDF extraction failed: {e}")
                return jsonify({"ok": True})
            # summarize & create LP
            title = "Extracted Lesson"
            summary = summarize_text(text, sentences_count=6)
            objectives = extract_objectives_from_text(text)
            activities = generate_activities(text)
            assessment = generate_assessment_questions(text)
            references = f"Source: PDF uploaded by user"
            fill_template_and_send(chat_id, title, summary, objectives, activities, assessment, references)
            return jsonify({"ok": True})

    # handle photo (OCR)
    if "photo" in msg:
        if not OCR_AVAILABLE:
            send_message(chat_id, "OCR not available. Please paste text or upload a PDF.")
            return jsonify({"ok": True})
        photo = msg["photo"][-1]
        tmpdir = tempfile.mkdtemp()
        local_path = os.path.join(tmpdir, "img.jpg")
        try:
            download_file(photo["file_id"], local_path)
            text = pytesseract.image_to_string(Image.open(local_path))
        except Exception as e:
            send_message(chat_id, f"OCR failed: {e}")
            return jsonify({"ok": True})
        summary = summarize_text(text, sentences_count=6)
        objectives = extract_objectives_from_text(text)
        activities = generate_activities(text)
        assessment = generate_assessment_questions(text)
        fill_template_and_send(chat_id, "Scanned Lesson", summary, objectives, activities, assessment, "Source: scanned image")
        return jsonify({"ok": True})

    # If user chooses "Paste Text" or pastes text
    if "text" in msg:
        txt = msg["text"].strip()
        # flow choices
        if txt == "Paste Text":
            send_message(chat_id, "Please paste the chapter text now.")
            SESS[chat_id]["state"] = "await_text"
            return jsonify({"ok": True})
        if txt == "Ask Bot to Find Lesson":
            SESS[chat_id]["state"] = "await_grade"
            send_message(chat_id, "Which Grade? (e.g., Grade 6)")
            return jsonify({"ok": True})
        # handle stateful conversation
        state = SESS[chat_id].get("state")
        if state == "await_text":
            text = txt
            summary = summarize_text(text, sentences_count=6)
            objectives = extract_objectives_from_text(text)
            activities = generate_activities(text)
            assessment = generate_assessment_questions(text)
            fill_template_and_send(chat_id, "Pasted Lesson", summary, objectives, activities, assessment, "Source: pasted by user")
            SESS[chat_id]["state"] = "idle"
            return jsonify({"ok": True})
        if state == "await_grade":
            SESS[chat_id]["tmp"] = {"grade": txt}
            SESS[chat_id]["state"] = "await_subject"
            send_message(chat_id, "Which Subject? (e.g., Mathematics)")
            return jsonify({"ok": True})
        if state == "await_subject":
            SESS[chat_id]["tmp"]["subject"] = txt
            SESS[chat_id]["state"] = "await_chapter"
            send_message(chat_id, "Which Chapter name or number?")
            return jsonify({"ok": True})
        if state == "await_chapter":
            SESS[chat_id]["tmp"]["chapter"] = txt
            # perform web search using DuckDuckGo
            grade = SESS[chat_id]["tmp"].get("grade")
            subject = SESS[chat_id]["tmp"].get("subject")
            chapter = SESS[chat_id]["tmp"].get("chapter")
            query = f"{grade} {subject} {chapter} summary lesson"
            send_message(chat_id, f"Searching web for: {query}")
            try:
                hits = ddg(query, max_results=5)
            except Exception as e:
                send_message(chat_id, f"Search failed: {e}. Proceeding without web sources.")
                hits = []
            # fetch and combine top page texts
            combined_texts = []
            references = []
            for h in hits[:3]:
                url = h.get("href") if "href" in h else h.get("url")
                title = h.get("title")
                snippet = h.get("body") or h.get("snippet","")
                txt = extract_text_from_url(url) if url else snippet
                if txt:
                    combined_texts.append(txt)
                references.append(f"{title if title else url} — {url}")
            big_text = "\n\n".join(combined_texts) or " ".join([h.get("body","") or h.get("snippet","") for h in hits])
            summary = summarize_text(big_text or chapter, sentences_count=6)
            objectives = extract_objectives_from_text(big_text or chapter)
            activities = generate_activities(big_text or chapter)
            assessment = generate_assessment_questions(big_text or chapter)
            ref_text = "\n".join(references) if references else "No web refs found"
            fill_template_and_send(chat_id, f"{chapter} ({subject})", summary, objectives, activities, assessment, ref_text)
            SESS[chat_id]["state"] = "idle"
            return jsonify({"ok": True})

        # if not in a flow and not a special command, ignore or prompt
        send_message(chat_id, "Send /Hi_Rise to begin or paste the chapter text.")
    return jsonify({"ok": True})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
